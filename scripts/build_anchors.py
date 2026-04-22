"""
build_anchors.py — Extract domain-grounded anchor sets for the MAVEN ML backend.

Reads:
  - perinatal_knowledge_base/dept_mch_feed/perinatal_misinformation.docx
  - perinatal_knowledge_base/dept_mch_feed/perinatal_comprehensive.docx

Writes (under maven_app/anchors/):
  - authority_anchors.json          list[str]
  - misinfo_anchors.json            list[{claim, evidence, domain}]
  - misinfo_type_anchors.json       dict[type_id -> list[str]]   (hand-curated seeds)
  - comprehensive_paragraphs.json   list[str]   (IsolationForest training corpus)

Idempotent. Re-run any time the source .docx files change.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from typing import Optional

from docx import Document
from docx.text.paragraph import Paragraph

# Paths
ROOT       = Path(__file__).resolve().parents[1]
FEED_DIR   = ROOT / 'perinatal_knowledge_base' / 'dept_mch_feed'
MISINFO_DOCX  = FEED_DIR / 'perinatal_misinformation.docx'
COMPRE_DOCX   = FEED_DIR / 'perinatal_comprehensive.docx'
OUT_DIR    = ROOT / 'maven_app' / 'anchors'

# Regex
GUIDELINE_RE = re.compile(
    r'\b(ACOG|ACNM|AWHONN|AAP|SMFM|SFP|NAF|USPSTF|CDC|AHRQ|HRSA|WHO|'
    r'MedlinePlus|NIH|NICHD|FDA|UNICEF|ACIP|LactMed|Cochrane|IOM|NAS)\b'
)
SECTION_LETTER_RE = re.compile(r'^(\d+)([A-Z])\.')

SECTION_TO_DOMAIN = {
    '1A': 'antenatal',
    '1B': 'intrapartum',
    '1C': 'fourth_trimester',
    '1D': 'family_planning',
    '2A': 'clinical_provider',
    '2B': 'clinical_gaslighting',
    '2C': 'clinical_guidelines',
    '3A': 'social_motherhood',
    '3B': 'social_family',
    '3C': 'social_media',
    '3D': 'social_help_seeking',
    '4A': 'info_structural',
    '4B': 'info_topic_misinfo',
    '4C': 'info_spread_mechanisms',
    '4D': 'info_correction_strategies',
    '4E': 'info_quality_signals',
}

# Hand-curated seed anchors for the JGIM 8-category misinfo taxonomy
# (John & Gorman, JGIM 2024, Table 2 — narrative review of reproductive health misinfo).
MISINFO_TYPE_SEEDS: dict[str, list[str]] = {
    'unattributed_risk': [
        'There are hidden long-term risks to this procedure that doctors will not tell you about.',
        'Many women suffer side effects from this medication but the data is being suppressed.',
        'Studies have shown serious harms, but the sources are never named.',
        'Hospitals do not disclose the real complication rates of this intervention.',
        'Independent researchers warn of dangers the medical establishment ignores.',
        'This treatment causes problems that mainstream science refuses to investigate.',
    ],
    'not_aligned_with_guidelines': [
        'You should refuse the routine glucose tolerance test during pregnancy.',
        'Skip the Group B Strep screening at 36 weeks because it is unnecessary.',
        'There is no need to follow ACOG recommendations on prenatal care visits.',
        'Standard induction protocols at 41 weeks are outdated and should be ignored.',
        'Mainstream guidelines on breastfeeding cessation are wrong about medications.',
        'The recommended vaccine schedule during pregnancy is not based on real evidence.',
    ],
    'promotes_alternative_medicine': [
        'Replace prenatal vitamins with herbal teas and whole foods only.',
        'Use essential oils instead of medication to manage gestational diabetes.',
        'Homeopathic remedies are safer and more effective than insulin during pregnancy.',
        'Chiropractic adjustments can turn a breech baby and avoid the need for medical intervention.',
        'Castor oil and natural induction methods are preferable to Pitocin.',
        'Postpartum depression should be treated with adaptogens, not antidepressants.',
    ],
    'exaggerated_risk': [
        'Epidurals during labour cause permanent paralysis in many women.',
        'Ultrasounds during pregnancy damage the developing fetal brain.',
        'Cesarean sections always result in lifelong autoimmune disease in the child.',
        'Pitocin during labour causes severe long-term harm to the baby.',
        'Vitamin K injections at birth dramatically increase the risk of childhood leukemia.',
        'Hospital births put the mother and baby at far greater risk than home birth.',
    ],
    'discourages_evidence_based': [
        'Decline the Tdap vaccine offered during the third trimester.',
        'Refuse continuous fetal monitoring because it leads to unnecessary cesareans.',
        'Do not consent to delayed cord clamping unless you understand the hidden risks.',
        'Skip the postpartum depression screening because the questions are biased.',
        'Refuse hepatitis B vaccination of your newborn at the hospital.',
        'Decline the routine eye ointment given to newborns immediately after birth.',
    ],
    'undermines_medical_trust': [
        'OB-GYNs only recommend cesareans because they are paid more for surgical deliveries.',
        'Hospitals push induction for scheduling convenience, not patient safety.',
        'Doctors do not really know what causes preeclampsia and are guessing.',
        'The medical system treats pregnant women as profit centres, not patients.',
        'Pharmaceutical companies suppress research that contradicts their products.',
        'Mainstream maternity care is built on outdated patriarchal assumptions.',
    ],
    'inaccurate_biological_mechanism': [
        'Plan B causes abortion by detaching an implanted embryo from the uterus.',
        'IUDs work by repeatedly aborting fertilized eggs every month.',
        'Breastfeeding alone is reliable contraception for the entire first year postpartum.',
        'A baby cannot feel pain during labour because the nervous system is not developed.',
        'Vaccines bypass the placenta and inject toxins directly into the fetus.',
        'Hormonal contraceptives accumulate in the body and cause permanent infertility.',
    ],
    'other': [
        'Real mothers do not need pain relief during childbirth.',
        'A natural birth is the only way to truly bond with your baby.',
        'If you needed a cesarean you failed at giving birth.',
        'Modern motherhood has been ruined by overmedicalization.',
        'Trust your body and ignore the constant testing during pregnancy.',
    ],
}


# Helpers
def style_name(p: Paragraph) -> str:
    try:
        return p.style.name if p.style else ''
    except Exception:
        return ''


def is_heading(p: Paragraph, level: Optional[int] = None) -> bool:
    sn = style_name(p)
    if not sn.startswith('Heading'):
        return False
    if level is None:
        return True
    return sn == f'Heading {level}'


def normalize(s: str) -> str:
    return re.sub(r'\s+', ' ', s).strip()


# Misinfo doc parser
def parse_misinfo_pairs(doc_path: Path) -> list[dict]:
    doc = Document(str(doc_path))
    pairs: list[dict] = []
    current_h1: str = ''
    current_h2: str = ''
    current_section_letter: str = ''
    pending_claim: Optional[str] = None

    for p in doc.paragraphs:
        text = normalize(p.text)
        if not text:
            continue

        if is_heading(p, 1):
            current_h1 = text
            current_h2 = ''
            current_section_letter = ''
            pending_claim = None
            continue
        if is_heading(p, 2):
            current_h2 = text
            m = SECTION_LETTER_RE.match(text)
            current_section_letter = f'{m.group(1)}{m.group(2)}' if m else ''
            pending_claim = None
            continue
        if is_heading(p, 3):
            pending_claim = None
            continue

        # CLAIM line
        if text.upper().startswith('CLAIM:'):
            pending_claim = text[len('CLAIM:'):].strip().lstrip("'\"").rstrip("'\"")
            continue

        # EVIDENCE line — must follow a CLAIM
        if text.upper().startswith('EVIDENCE:') and pending_claim:
            evidence = text[len('EVIDENCE:'):].strip()
            domain = SECTION_TO_DOMAIN.get(
                current_section_letter,
                current_h2.lower().replace(' ', '_')[:40] or 'unknown',
            )
            pairs.append({
                'claim':    pending_claim,
                'evidence': evidence,
                'domain':   domain,
            })
            pending_claim = None

    return pairs


# Comprehensive doc parser
def extract_authority_sentences(doc_path: Path) -> list[str]:
    doc = Document(str(doc_path))
    out: list[str] = []
    seen: set[str] = set()
    in_part_one = False
    in_domain_one = False

    for p in doc.paragraphs:
        text = normalize(p.text)
        if not text:
            continue

        if is_heading(p, 1):
            heading = text.upper()
            in_part_one  = heading.startswith('PART I:') or 'AUTHORITATIVE GUIDELINE' in heading
            in_domain_one = heading.startswith('DOMAIN 1')
            continue

        # Restrict to PART I (authority bodies) and DOMAIN 1 (clinical care content) for high signal
        if not (in_part_one or in_domain_one):
            continue

        # Filter: must contain a guideline body, be a single declarative sentence
        if not GUIDELINE_RE.search(text):
            continue
        if len(text) < 50 or len(text) > 400:
            continue
        # Single-sentence-ish: at most one terminal period
        if text.count('. ') > 1:
            # Try to split and keep the first clause containing a guideline body
            clauses = [c.strip() for c in re.split(r'(?<=[.;])\s+', text) if c.strip()]
            for c in clauses:
                if GUIDELINE_RE.search(c) and 50 <= len(c) <= 300 and c not in seen:
                    out.append(c)
                    seen.add(c)
            continue

        if text not in seen:
            out.append(text)
            seen.add(text)

    return out


def extract_comprehensive_paragraphs(doc_path: Path) -> list[str]:
    """Every non-heading paragraph >40 chars from the comprehensive doc.
    Used as the IsolationForest training corpus — the model learns the
    distribution of authoritative perinatal-text language.
    """
    doc = Document(str(doc_path))
    paras: list[str] = []
    seen: set[str] = set()
    for p in doc.paragraphs:
        text = normalize(p.text)
        if not text or len(text) < 40:
            continue
        if is_heading(p):
            continue
        if text in seen:
            continue
        seen.add(text)
        paras.append(text)
    return paras


# Main
def main() -> int:
    sys.stdout.reconfigure(encoding='utf-8')

    if not MISINFO_DOCX.exists():
        print(f'[ERROR] missing: {MISINFO_DOCX}', file=sys.stderr)
        return 1
    if not COMPRE_DOCX.exists():
        print(f'[ERROR] missing: {COMPRE_DOCX}', file=sys.stderr)
        return 1

    OUT_DIR.mkdir(parents=True, exist_ok=True)

    print('[build_anchors] Parsing misinformation doc...')
    pairs = parse_misinfo_pairs(MISINFO_DOCX)
    print(f'  -> {len(pairs)} CLAIM/EVIDENCE pairs')

    domain_counts: dict[str, int] = {}
    for pair in pairs:
        domain_counts[pair['domain']] = domain_counts.get(pair['domain'], 0) + 1
    for dom, n in sorted(domain_counts.items(), key=lambda kv: -kv[1]):
        print(f'     {dom:25s} {n}')

    print('\n[build_anchors] Extracting authority sentences from comprehensive doc...')
    auth_from_compre = extract_authority_sentences(COMPRE_DOCX)
    print(f'  -> {len(auth_from_compre)} sentences with guideline body keyword')

    # Also harvest evidence lines from misinfo pairs as authority anchors
    auth_from_evidence = []
    seen_auth: set[str] = set()
    for pair in pairs:
        ev = pair['evidence']
        # Take leading clause(s) up to ~280 chars to keep anchors atomic
        first = re.split(r'(?<=[.;])\s+', ev)[0].strip()
        if 40 <= len(first) <= 300 and first not in seen_auth:
            auth_from_evidence.append(first)
            seen_auth.add(first)

    authority_anchors = list(dict.fromkeys(auth_from_compre + auth_from_evidence))
    print(f'  -> {len(authority_anchors)} unique authority anchors total')

    print('\n[build_anchors] Misinfo type seed counts:')
    for tid, seeds in MISINFO_TYPE_SEEDS.items():
        print(f'     {tid:35s} {len(seeds)}')

    print('\n[build_anchors] Extracting comprehensive paragraphs (IsoForest corpus)...')
    comp_paragraphs = extract_comprehensive_paragraphs(COMPRE_DOCX)
    print(f'  -> {len(comp_paragraphs)} paragraphs (>40 chars, non-heading)')

    # Write JSON
    (OUT_DIR / 'authority_anchors.json').write_text(
        json.dumps(authority_anchors, indent=2, ensure_ascii=False), encoding='utf-8'
    )
    (OUT_DIR / 'misinfo_anchors.json').write_text(
        json.dumps(pairs, indent=2, ensure_ascii=False), encoding='utf-8'
    )
    (OUT_DIR / 'misinfo_type_anchors.json').write_text(
        json.dumps(MISINFO_TYPE_SEEDS, indent=2, ensure_ascii=False), encoding='utf-8'
    )
    (OUT_DIR / 'comprehensive_paragraphs.json').write_text(
        json.dumps(comp_paragraphs, indent=2, ensure_ascii=False), encoding='utf-8'
    )

    print(f'\n[build_anchors] Wrote 4 files to {OUT_DIR}')
    print('  - authority_anchors.json')
    print('  - misinfo_anchors.json')
    print('  - misinfo_type_anchors.json')
    print('  - comprehensive_paragraphs.json')
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
